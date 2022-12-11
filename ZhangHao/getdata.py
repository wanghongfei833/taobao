# -*- coding: utf-8 -*-
# @Time    : 2022/12/6 13:03
# @Author  : HongFei Wang
import os
import sys
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt, RGBColor
import docx
import tqdm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 设置对象居中、对齐等。
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER  # 设置制表符等
from docx.shared import Inches  # 设置图像大小
from docx.shared import Pt  # 设置像素、缩进等
from docx.shared import RGBColor  # 设置字体颜色
from docx.shared import Length  # 设置宽度


def hebing(table, *args):
    cell1 = table.cell(args[0], args[1])
    cell2 = table.cell(args[2], args[3])
    cell1.merge(cell2)


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "FF0000", "space": "0"},
        bottom={"sz": 12, "color": "00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        # edge_data = {"sz": 6,
        #              "color": "#000000",
        #              "val": "single"}
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def write(table_file, table_save, row, col):
    table_save.cell(row, col).text = table_file.cell(row, col).text.strip('\n')
    table_save.cell(row, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_save.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 水平居中
    for par in table_save.cell(row, col).paragraphs:
        ##horizontal centered
        par.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  #
        for run in par.runs:
            run.font.size = Pt(12)
            r = run._element
            run.font.name = '宋体'
            r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    set_cell_border(table_save.cell(row, col),
                    top={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                    bottom={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                    left={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                    right={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                    insideH={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                    end={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"})


def write2(texts, table_save, row, col):
    table_save.cell(row, col).text = texts
    table_save.cell(row, col).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table_save.cell(row, col).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 水平居中
    for par in table_save.cell(row, col).paragraphs:
        ##horizontal centered
        par.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  #
        for run in par.runs:
            run.font.size = Pt(12)
            r = run._element
            run.font.name = '宋体'
            r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    set_cell_border(table_save.cell(row, col),
                    top={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                    bottom={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                    left={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                    right={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                    insideH={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"},
                    end={"sz": 0.5, "val": "single", "color": "#000000", "space": "0"})


doc_Save = docx.Document()


def add_word(file):
    doc = docx.Document(file)
    # 第一页 宋体 四号
    text_par = doc.paragraphs[2].text
    paragraph = doc_Save.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph = paragraph.add_run(text_par)
    paragraph.font.name = u'宋体'
    paragraph.bold = True
    # 设置中文字体
    paragraph._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 设置字体大小
    paragraph.font.size = Pt(14)

    # 第二页 宋体 小四号
    text_par = doc.paragraphs[3].text
    paragraph = doc_Save.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph = paragraph.add_run(text_par)
    paragraph.font.name = u'宋体'
    # 设置中文字体
    paragraph._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 设置字体大小
    paragraph.font.size = Pt(12)

    table_file = doc.tables[0]
    rows, cols = len(table_file.rows), len(table_file.columns)
    rows = 18
    all_list = []

    table_save = doc_Save.add_table(18, cols)
    table_save.alignment = WD_TABLE_ALIGNMENT.CENTER
    #
    # for row in range(rows):
    #     for col in range(cols):
    #         cell = table_save.cell(row, col)
    #         # set_cell_border(cell=cell,
    #         #                 top={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"},
    #         #                 bottom={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"},
    #         #                 left={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"},
    #         #                 right={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"},
    #         #                 insideH={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"},
    #         #                 end={"sz": 0.5, "val": "double", "color": "#000000", "space": "0"})
    #
    #         cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中
    #         cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 垂直居中
    #         run = cell.paragraphs[0].add_run(u'')
    #         run.font.name = '宋体'
    #         run.font.size = 12
    for row in range(rows):
        if row <= 13:
            # 前15行
            # table_save.cell(row, 0).text = table_file.cell(row, 0).text
            write(table_file, table_save, row, 0)
            hebing(table_save, row, 0, row, 1)
            # table_save.cell(row, 2).text = table_file.cell(row, 2).text
            write(table_file, table_save, row, 2)
            hebing(table_save, row, 2, row, cols - 1)
        elif row == 14:
            hebing(table_save, row, 0, row, cols - 1)
            # table_save.cell(row, 0).text = table_file.cell(row, 0).text
            write(table_file, table_save, row, 0)

        elif row > 14:
            for col in range(cols):
                try:
                    all_list.append(table_file.cell(row, col).text.replace('\n', ''))
                except:
                    pass
    all_list = ' '.join(all_list).split(' ')
    for row in range(15, rows):
        if row > 18: break
        name = all_list[(row - 15) * 6]
        xingbie = all_list[(row - 15) * 6 + 1]
        guanxi = all_list[(row - 15) * 6 + 3]
        id_Card = all_list[(row - 15) * 6 + 4]
        beizhu = all_list[(row - 15) * 6 + 5]
        write2(name, table_save, row, 0)
        write2(xingbie, table_save, row, 1)
        write2(guanxi, table_save, row, 3)
        write2(id_Card, table_save, row, 4)
        write2(beizhu, table_save, row, 5)

        # table_save.cell(row, 0).text = name
        # table_save.cell(row, 1).text = xingbie
        # table_save.cell(row, 3).text = guanxi
        # table_save.cell(row, 4).text = id_Card
        # table_save.cell(row, 5).text = beizhu

        hebing(table_save, row, 1, row, 2)
        hebing(table_save, row, 5, row, 6)


def make_docx(file_path, save_path):
    file = file_path
    makes_par = tqdm.tqdm(os.listdir(file))
    for dirs in makes_par:
        files = os.path.join(file, dirs)
        add_word(files)
        makes_par.set_description_str('生成word中')

    doc_Save.save(save_path)
