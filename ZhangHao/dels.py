# -*- coding: utf-8 -*-
# @Time    : 2022/12/6 23:53
# @Author  : HongFei Wang
import sys

import docx
import os

import tqdm
from docx import Document
import docxcompose.composer as coms



# 删除指定段落
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None


# 定义合并文档的函数
def merge_doc(source_file_path_list, target_file_path):
    """
    :param source_file_path_list: 源文件路径列表
    :param target_file_path: 目标文件路径
    """
    # 填充分页符号文档
    page_break_doc = Document()
    page_break_doc.add_page_break()
    # 定义新文档
    target_doc = Document(source_file_path_list[0])
    target_composer = coms.Composer(target_doc)
    print('第一步')
    for i in range(len(source_file_path_list)):
        # 跳过第一个作为模板的文件
        if i == 0:
            continue
        try:
            target_composer.append(page_break_doc)
        except Exception as err:
            print(err)
            print('出错')
            return
        # 拼接文档内容
        f = source_file_path_list[i]
        target_composer.append(Document(f))
    # 保存目标文档
    target_composer.save(target_file_path)


def run(file, save):
    temp_file = file
    for sss in os.listdir(temp_file):
        document = docx.Document(os.path.join(temp_file, sss))
        tables = document.tables

        for index, paragraph in enumerate(document.paragraphs):
            if index == 2 or index == 3: continue
            p = paragraph._element
            p.getparent().remove(p)
            paragraph._p = paragraph._element = None

        for index, table in enumerate(tables):
            if index == 0: continue
            table._element.getparent().remove(table._element)
        document.save(os.path.join(file, sss))
    # 原文件夹绝对路径, python能自动找到当前文件夹在当前电脑上的绝对路径
    source_path = file
    # 目标文件路径
    target_file = save
    source_file_list = tqdm.tqdm(os.listdir(source_path))
    new_list = []
    for item in source_file_list:
        if item.endswith('.docx'):
            new_list.append(item)
        source_file_list.set_description_str('word生成中')
    # 获取源文件夹内内文件列表
    source_file_list_all = []
    for file in new_list:
        source_file_list_all.append(source_path + '\\' + file)
    # 调用合并文档函数
    merge_doc(source_file_list_all, target_file)
    print('删除临时文件夹')
    for path in os.listdir(temp_file):
        os.remove(os.path.join(temp_file,path))
    os.rmdir(source_path)
    input('输入任意键结束程序:...')
    sys.exit()